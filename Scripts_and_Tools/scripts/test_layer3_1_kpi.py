from docx import Document
import os

os.makedirs(r"D:\Project_Hub\02_去識別化數據", exist_ok=True)

# 建立 KPI 標案需求書
rfp = Document()
rfp.add_heading('114年度標案需求書', 0)
rfp.add_heading('肆、績效指標 (KPI)', 1)
rfp.add_paragraph('本專案年度量化目標：體驗人次需達 10,000 人次。')
rfp.add_paragraph('滿意度必須達到 85% 以上。')
rfp.add_paragraph('辦理場次至少 10 場。')
rfp.save(r"D:\Project_Hub\02_去識別化數據\RFP_KPI_TARGETS.docx")

# 建立 KPI 計畫書 (部分未達標)
plan = Document()
plan.add_heading('114年度計畫書', 0)
plan.add_heading('陸、績效指標', 1)
plan.add_paragraph('本計畫預計年度體驗人次可達 8,000 人次。') # 未達標
plan.add_paragraph('目標場次為 12 場。') # 達標
plan.add_paragraph('滿意度預期目標為 90%。') # 達標
plan.save(r"D:\Project_Hub\02_去識別化數據\PLAN_KPI_PERFORMANCE.docx")
