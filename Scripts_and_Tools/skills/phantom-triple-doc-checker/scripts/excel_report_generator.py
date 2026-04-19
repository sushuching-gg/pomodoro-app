import argparse
import os
import sys
import re
from datetime import datetime

if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)

try:
    from docx import Document
except ImportError:
    print("\u932f\u8aa4: \u672a\u5b89\u88dd python-docx\u3002\u8acb\u57f7\u884c: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    print("\u932f\u8aa4: \u672a\u5b89\u88dd openpyxl\u3002\u8acb\u57f7\u884c: pip install openpyxl")
    sys.exit(1)

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import fitz  # PyMuPDF
    import easyocr
    import numpy as np
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# ============================================================
# Constants
# ============================================================

COMPARE_ITEMS = [
    {
        'category': '\u8a08\u756b\u57fa\u672c\u8cc7\u6599',
        'items': [
            {'name': '\u8a08\u756b\u540d\u7a31', 'keywords': ['\u8a08\u756b\u540d\u7a31', '\u5c08\u6848\u540d\u7a31']},
            {'name': '\u8a08\u756b\u671f\u7a0b', 'keywords': ['\u671f\u7a0b', '\u57f7\u884c\u671f\u9593', '\u5c65\u7d04\u671f']},
            {'name': '\u8a08\u756b\u7d93\u8cbb', 'keywords': ['\u7d93\u8cbb', '\u9810\u7b97', '\u91d1\u984d', '\u5408\u7d04\u50f9\u91d1', '\u7e3d\u8a08', '\u5c0f\u8a08']},
            {'name': '\u57f7\u884c\u55ae\u4f4d', 'keywords': ['\u57f7\u884c\u55ae\u4f4d', '\u59d4\u8a17', '\u5ee0\u5546', '\u5f97\u6a19']},
        ]
    },
    {
        'category': '\u5de5\u4f5c\u7bc4\u570d',
        'items': [
            {'name': '\u4e3b\u8981\u5de5\u4f5c\u9805\u76ee', 'keywords': ['\u5de5\u4f5c\u9805\u76ee', '\u670d\u52d9\u9805\u76ee', '\u57f7\u884c\u5167\u5bb9']},
            {'name': '\u7cfb\u7d71\u958b\u767c', 'keywords': ['\u7cfb\u7d71\u958b\u767c', '\u7a0b\u5f0f\u958b\u767c', '\u8edf\u9ad4\u958b\u767c', 'APP', '\u7db2\u7ad9']},
            {'name': '\u670d\u52d9\u7bc4\u570d', 'keywords': ['\u670d\u52d9\u7bc4\u570d', '\u670d\u52d9\u5c0d\u8c61', '\u4f7f\u7528\u8005']},
        ]
    },
    {
        'category': '\u6280\u8853\u898f\u683c',
        'items': [
            {'name': '\u7cfb\u7d71\u67b6\u69cb', 'keywords': ['\u7cfb\u7d71\u67b6\u69cb', '\u6280\u8853\u67b6\u69cb', '\u96f2\u7aef', '\u4f3a\u670d\u5668']},
            {'name': '\u8cc7\u5b89\u9632\u8b77', 'keywords': ['\u8cc7\u5b89', '\u8cc7\u8a0a\u5b89\u5168', '\u5b89\u5168\u6a5f\u5236', 'HTTPS', 'SSL']},
            {'name': '\u7cfb\u7d71\u6574\u5408', 'keywords': ['\u6574\u5408', '\u4e32\u63a5', 'API', 'Google Map', '\u53f0\u5317\u901a', '\u91d1\u6d41']},
        ]
    },
    {
        'category': '\u4eba\u529b\u914d\u7f6e',
        'items': [
            {'name': '\u5718\u968a\u7d44\u6210', 'keywords': ['\u5718\u968a', '\u4eba\u54e1', '\u7d44\u7e54', '\u5c08\u6848\u7d93\u7406']},
            {'name': '\u99d0\u9ede\u4eba\u54e1', 'keywords': ['\u99d0\u9ede', '\u73fe\u5834', '\u652f\u63f4']},
        ]
    },
    {
        'category': '\u6642\u7a0b\u898f\u5283',
        'items': [
            {'name': '\u6574\u9ad4\u6642\u7a0b', 'keywords': ['\u6642\u7a0b', '\u7518\u7279\u5716', '\u91cc\u7a0b\u7891', '\u968e\u6bb5']},
            {'name': '\u4ea4\u4ed8\u9805\u76ee', 'keywords': ['\u4ea4\u4ed8', '\u9a57\u6536', '\u6210\u679c\u5831\u544a']},
        ]
    },
    {
        'category': '\u7d93\u8cbb\u7de8\u5217',
        'items': [
            {'name': '\u7d93\u8cbb\u660e\u7d30', 'keywords': ['\u7d93\u8cbb', '\u9810\u7b97', '\u8cbb\u7528', '\u55ae\u50f9', '\u91d1\u984d']},
            {'name': '\u4ed8\u6b3e\u689d\u4ef6', 'keywords': ['\u4ed8\u6b3e', '\u8acb\u6b3e', '\u6491\u4ed8', '\u671f\u6b3e']},
        ]
    },
    {
        'category': 'KPI/\u7e3e\u6548\u6307\u6a19',
        'items': [
            {'name': '\u7e3e\u6548\u6307\u6a19', 'keywords': ['KPI', '\u7e3e\u6548', '\u6307\u6a19', '\u76ee\u6a19\u503c', '\u9054\u6210\u7387']},
            {'name': '\u6548\u76ca\u8a55\u4f30', 'keywords': ['\u6548\u76ca', '\u8a55\u4f30', '\u6210\u679c']},
        ]
    },
]


# ============================================================
# Document Reader
# ============================================================

def read_document(path):
    if not os.path.exists(path):
        return {'text': '', 'filename': os.path.basename(path), 'pages': 0}
    
    ext = os.path.splitext(path)[1].lower()
    filename = os.path.basename(path)
    
    if ext == '.docx':
        return _read_docx(path, filename)
    elif ext == '.pdf':
        return _read_pdf(path, filename)
    else:
        return {'text': '', 'filename': filename, 'pages': 0}

def _read_docx(path, filename):
    try:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_texts.append(' '.join(cells))
        
        full_text = '\n'.join(paragraphs + table_texts)
        return {'text': full_text, 'filename': filename, 'pages': len(doc.paragraphs)}
    except Exception as e:
        print(f'\u8b66\u544a: \u7121\u6cd5\u8b80\u53d6 {filename}: {e}')
        return {'text': '', 'filename': filename, 'pages': 0}

def _read_pdf(path, filename):
    if not HAS_PDF:
        print('\u8b66\u544a: \u672a\u5b89\u88dd PyPDF2\uff0c\u7121\u6cd5\u8b80\u53d6 PDF')
        return {'text': '', 'filename': filename, 'pages': 0}
    try:
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f'[p.{i+1}] {text}')
        full_text = '\n'.join(pages)
        return {'text': full_text, 'filename': filename, 'pages': len(reader.pages)}
    except Exception as e:
        print(f'\u8b66\u544a: \u7121\u6cd5\u8b80\u53d6 PDF {filename}: {e}')
        return {'text': '', 'filename': filename, 'pages': 0}


def find_relevant_content(doc_data, keywords, max_chars=300):
    text = doc_data['text']
    if not text:
        return '\u672a\u5075\u6e2c\u5230\u76f8\u95dc\u5167\u5bb9'
    
    results = []
    lines = text.split('\n')
    
    for line in lines:
        if any(kw in line for kw in keywords):
            clean = line.strip()
            # Try to find page reference
            page_match = re.search(r'\[p\.(\d+)\]', clean)
            page_ref = f'(p.{page_match.group(1)})' if page_match else ''
            clean = re.sub(r'\[p\.\d+\]\s*', '', clean)
            
            if len(clean) > 5:
                entry = clean[:150]
                if page_ref:
                    entry += f' {page_ref}'
                results.append(entry)
    
    if results:
        combined = '\n'.join(results[:5])
        if len(combined) > max_chars:
            combined = combined[:max_chars] + '...'
        return combined
    
    return '\u672a\u5075\u6e2c\u5230\u76f8\u95dc\u5167\u5bb9'


def judge_consistency(plan_content, rfp_content, proposal_content):
    has_plan = '\u672a\u5075\u6e2c' not in plan_content
    has_rfp = '\u672a\u5075\u6e2c' not in rfp_content
    has_proposal = '\u672a\u5075\u6e2c' not in proposal_content
    
    if has_plan and has_rfp and has_proposal:
        return '\u2713', '\u2713', '\u7b26\u5408'
    elif has_plan and has_proposal and not has_rfp:
        return '\u2713', '\u26a0', '\u5f85\u78ba\u8a8d'
    elif (has_plan or has_rfp) and not has_proposal:
        return '\u26a0', '\u26a0', '\u5f85\u8aaa\u660e'
    elif not has_plan and not has_rfp and not has_proposal:
        return '-', '-', '\u7121\u8cc7\u6599'
    else:
        return '\u26a0', '\u2713', '\u5f85\u78ba\u8a8d'


# ============================================================
# Excel Generator
# ============================================================

def generate_excel(plan_data, rfp_data, proposal_data, output_path, city=None, year=None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '\u4e09\u66f8\u6bd4\u5c0d\u8868'
    
    # --- Styles ---
    title_font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=14, bold=True)
    header_font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=10, bold=True, color='FFFFFF')
    normal_font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=9)
    category_font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=10, bold=True)
    
    header_fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    category_fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
    ok_fill = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
    warn_fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    bad_fill = PatternFill(start_color='FCE4EC', end_color='FCE4EC', fill_type='solid')
    
    wrap_align = Alignment(wrap_text=True, vertical='top')
    center_align = Alignment(wrap_text=True, vertical='center', horizontal='center')
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # --- Column widths ---
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 45
    ws.column_dimensions['D'].width = 40
    ws.column_dimensions['E'].width = 45
    ws.column_dimensions['F'].width = 8
    ws.column_dimensions['G'].width = 8
    ws.column_dimensions['H'].width = 10
    ws.column_dimensions['I'].width = 55
    
    # === Row 1: Title ===
    title_text = ''
    if city:
        title_text = f'{city}'
    if year:
        title_text = f'{year}\u5e74\u5ea6 {title_text}'
    title_text += '\u8a08\u756b\u4e09\u66f8\u6bd4\u5c0d\u8868'
    
    ws.merge_cells('A1:I1')
    title_cell = ws['A1']
    title_cell.value = title_text
    title_cell.font = title_font
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 35
    
    # === Row 2: Sub-info ===
    ws.merge_cells('A2:I2')
    info_cell = ws['A2']
    info_cell.value = f'\u7522\u751f\u6642\u9593: {datetime.now().strftime("%Y-%m-%d %H:%M")}  |  \u6838\u5b9a\u8a08\u756b\u66f8: {plan_data["filename"]}  |  \u62db\u6a19\u9700\u6c42\u66f8: {rfp_data["filename"]}  |  \u670d\u52d9\u5efa\u8b70\u66f8: {proposal_data["filename"]}'
    info_cell.font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=8, color='666666')
    info_cell.alignment = Alignment(wrap_text=True)
    ws.row_dimensions[2].height = 30
    
    # === Row 3: Headers ===
    headers = ['\u9805\u76ee\u5927\u985e', '\u6aa2\u67e5\u9805\u76ee', '\u6838\u5b9a\u8a08\u756b\u66f8', '\u62db\u6a19\u9700\u6c42\u66f8', '\u5ee0\u5546\u63d0\u6848(\u670d\u52d9\u5efa\u8b70\u66f8)', '\u627f\u63a5\u6027', '\u53ef\u884c\u6027', '\u7d9c\u5408\u5224\u5b9a', '\u6bd4\u5c0d\u8aaa\u660e\u8207\u5efa\u8b70']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border
    ws.row_dimensions[3].height = 30
    
    # === Data rows ===
    current_row = 4
    diff_count = 0
    
    for cat in COMPARE_ITEMS:
        category = cat['category']
        items = cat['items']
        
        for idx, item in enumerate(items):
            row = current_row
            
            # Column A: Category
            cell_a = ws.cell(row=row, column=1, value=category if idx == 0 else '')
            cell_a.font = category_font if idx == 0 else normal_font
            if idx == 0:
                cell_a.fill = category_fill
            cell_a.alignment = wrap_align
            cell_a.border = thin_border
            
            # Column B: Item name
            cell_b = ws.cell(row=row, column=2, value=item['name'])
            cell_b.font = normal_font
            cell_b.alignment = wrap_align
            cell_b.border = thin_border
            
            # Columns C/D/E: Document content
            plan_content = find_relevant_content(plan_data, item['keywords'])
            rfp_content = find_relevant_content(rfp_data, item['keywords'])
            proposal_content = find_relevant_content(proposal_data, item['keywords'])
            
            for col, content in [(3, plan_content), (4, rfp_content), (5, proposal_content)]:
                cell = ws.cell(row=row, column=col, value=content)
                cell.font = normal_font
                cell.alignment = wrap_align
                cell.border = thin_border
            
            # Columns F/G/H: Judgment
            continuity, feasibility, verdict = judge_consistency(plan_content, rfp_content, proposal_content)
            
            cell_f = ws.cell(row=row, column=6, value=continuity)
            cell_f.font = normal_font
            cell_f.alignment = center_align
            cell_f.border = thin_border
            
            cell_g = ws.cell(row=row, column=7, value=feasibility)
            cell_g.font = normal_font
            cell_g.alignment = center_align
            cell_g.border = thin_border
            
            cell_h = ws.cell(row=row, column=8, value=verdict)
            cell_h.font = normal_font
            cell_h.alignment = center_align
            cell_h.border = thin_border
            
            # Color-code verdict
            if verdict == '\u7b26\u5408':
                cell_h.fill = ok_fill
            elif verdict == '\u5f85\u78ba\u8a8d' or verdict == '\u5f85\u8aaa\u660e':
                cell_h.fill = warn_fill
                diff_count += 1
            elif verdict == '\u7121\u8cc7\u6599':
                cell_h.fill = bad_fill
                diff_count += 1
            
            # Column I: Notes
            note = _generate_note(plan_content, rfp_content, proposal_content, verdict)
            cell_i = ws.cell(row=row, column=9, value=note)
            cell_i.font = normal_font
            cell_i.alignment = wrap_align
            cell_i.border = thin_border
            
            # Row height
            ws.row_dimensions[row].height = 60
            current_row += 1
    
    # === Summary row ===
    current_row += 1
    ws.merge_cells(f'A{current_row}:I{current_row}')
    summary_cell = ws.cell(row=current_row, column=1, value='\u6574\u9ad4\u6bd4\u5c0d\u7d50\u8ad6')
    summary_cell.font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=12, bold=True)
    summary_cell.fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    summary_cell.font = Font(name='\u5fae\u8edf\u6b63\u9ed1\u9ad4', size=12, bold=True, color='FFFFFF')
    summary_cell.border = thin_border
    
    current_row += 1
    ws.merge_cells(f'A{current_row}:I{current_row}')
    total_items = sum(len(c['items']) for c in COMPARE_ITEMS)
    ok_items = total_items - diff_count
    
    conclusion = f'\u5171\u6bd4\u5c0d {total_items} \u9805\uff0c\u5176\u4e2d {ok_items} \u9805\u7b26\u5408\u3001{diff_count} \u9805\u5f85\u78ba\u8a8d/\u5f85\u8aaa\u660e\u3002\n'
    if diff_count == 0:
        conclusion += '\u4e09\u66f8\u5167\u5bb9\u4e00\u81f4\u6027\u826f\u597d\uff0c\u5efa\u8b70\u53ef\u4ee5\u9032\u884c\u8a2a\u8996\u3002'
    else:
        conclusion += f'\u8acb\u91dd\u5c0d {diff_count} \u9805\u5dee\u7570\u9032\u884c\u4eba\u5de5\u78ba\u8a8d\u3002'
    conclusion += f'\n\n\u6b64\u5831\u544a\u7531 Phantom \u4e09\u66f8\u6bd4\u5c0d Agent \u81ea\u52d5\u7522\u751f\uff0c\u50c5\u4f9b\u53c3\u8003\u3002'
    
    summary_detail = ws.cell(row=current_row, column=1, value=conclusion)
    summary_detail.font = normal_font
    summary_detail.alignment = wrap_align
    summary_detail.border = thin_border
    ws.row_dimensions[current_row].height = 70
    
    # === Save ===
    wb.save(output_path)
    print(f'\u2713 Excel \u4e09\u66f8\u6bd4\u5c0d\u8868\u5df2\u5132\u5b58\u81f3: {output_path}')


def _generate_note(plan, rfp, proposal, verdict):
    has_p = '\u672a\u5075\u6e2c' not in plan
    has_r = '\u672a\u5075\u6e2c' not in rfp
    has_pr = '\u672a\u5075\u6e2c' not in proposal
    
    if verdict == '\u7b26\u5408':
        return '\u4e09\u66f8\u5747\u5305\u542b\u76f8\u95dc\u5167\u5bb9\uff0c\u4e00\u81f4\u6027\u826f\u597d\u3002'
    elif verdict == '\u7121\u8cc7\u6599':
        return '\u4e09\u66f8\u5747\u672a\u5075\u6e2c\u5230\u76f8\u95dc\u5167\u5bb9\uff0c\u8acb\u4eba\u5de5\u78ba\u8a8d\u3002'
    else:
        missing = []
        if not has_p: missing.append('\u6838\u5b9a\u8a08\u756b\u66f8')
        if not has_r: missing.append('\u62db\u6a19\u9700\u6c42\u66f8')
        if not has_pr: missing.append('\u670d\u52d9\u5efa\u8b70\u66f8')
        return f'\u4ee5\u4e0b\u6587\u4ef6\u672a\u5075\u6e2c\u5230\u76f8\u95dc\u5167\u5bb9: {", ".join(missing)}\u3002\u8acb\u78ba\u8a8d\u662f\u5426\u5728\u5176\u4ed6\u7ae0\u7bc0\u6216\u9644\u4ef6\u4e2d\u3002'


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='Phantom \u4e09\u66f8\u6bd4\u5c0d Excel \u7522\u751f\u5668')
    parser.add_argument('--plan', required=True, help='\u6838\u5b9a\u8a08\u756b\u66f8\u8def\u5f91 (.docx/.pdf)')
    parser.add_argument('--rfp', required=True, help='\u62db\u6a19\u9700\u6c42\u66f8\u8def\u5f91 (.docx/.pdf)')
    parser.add_argument('--proposal', required=True, help='\u670d\u52d9\u5efa\u8b70\u66f8\u8def\u5f91 (.docx/.pdf)')
    parser.add_argument('--output', required=True, help='\u8f38\u51fa Excel \u6a94\u6848\u8def\u5f91 (.xlsx)')
    parser.add_argument('--city', default=None, help='\u7e23\u5e02\u540d\u7a31 (\u9078\u586b)')
    parser.add_argument('--year', default=None, help='\u5e74\u5ea6 (\u9078\u586b)')

    args = parser.parse_args()

    print('=' * 50)
    print('  Phantom \u4e09\u66f8\u6bd4\u5c0d Excel \u7522\u751f\u5668')
    print('=' * 50)
    print()

    print('[\u6b65\u9a5f 1/4] \u8b80\u53d6\u6838\u5b9a\u8a08\u756b\u66f8...')
    plan_data = read_document(args.plan)
    print(f'  \u2192 {plan_data["filename"]} ({plan_data["pages"]} \u9801/\u6bb5)')

    print('[\u6b65\u9a5f 2/4] \u8b80\u53d6\u62db\u6a19\u9700\u6c42\u66f8...')
    rfp_data = read_document(args.rfp)
    print(f'  \u2192 {rfp_data["filename"]} ({rfp_data["pages"]} \u9801/\u6bb5)')

    print('[\u6b65\u9a5f 3/4] \u8b80\u53d6\u670d\u52d9\u5efa\u8b70\u66f8...')
    proposal_data = read_document(args.proposal)
    print(f'  \u2192 {proposal_data["filename"]} ({proposal_data["pages"]} \u9801/\u6bb5)')

    print('[\u6b65\u9a5f 4/4] \u7522\u751f Excel \u6bd4\u5c0d\u8868...')
    generate_excel(plan_data, rfp_data, proposal_data, args.output, args.city, args.year)
    print('\n\u6bd4\u5c0d\u5b8c\u6210\u3002')


if __name__ == '__main__':
    main()
