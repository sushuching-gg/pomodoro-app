import sys
import docx

def analyze_docx(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"=== Document Analysis: {file_path} ===")
        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        print(f"Total Tables: {len(doc.tables)}\n")
        
        print("--- Headings & Key Text (First 500 chars of relevant paras) ---")
        count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Print headings or potential key fields
            if para.style.name.startswith('Heading') or len(text) < 50 or "經費" in text or "指標" in text or "KPI" in text:
                print(f"[{para.style.name}] {text[:60]}...")
                count += 1
                if count > 30: # Limit output
                    break
                    
        print("\n--- Table Analysis ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows")
            # Print first row cells to guess table content
            if table.rows:
                row_text = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"  Header: {row_text}")
            if i >= 5: # Limit table output
                break

    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        analyze_docx(sys.argv[1])
    else:
        print("Usage: python analyze_docx.py <path>")
