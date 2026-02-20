import argparse
import os
import sys

# Try to import dependencies
try:
    from docx import Document
except ImportError:
    print("錯誤: 未安裝 python-docx。請執行: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("錯誤: 未安裝 openpyxl。請執行: pip install openpyxl")
    sys.exit(1)

def extract_plan_text(docx_path):
    """
    Extracts text from a policy plan (Word document).
    """
    if not os.path.exists(docx_path):
        return f"錯誤: 找不到檔案: {docx_path}"

    try:
        doc = Document(docx_path)
        full_text = []
        
        # Keywords for section identification
        keywords = ["計畫目標", "工作項目", "KPI", "關鍵績效指標", "Goal", "Objective"]
        found_sections = {}
        
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for keywords
            is_header = False
            for kw in keywords:
                if kw in text:
                    current_section = kw
                    if current_section not in found_sections:
                        found_sections[current_section] = []
                    is_header = True
                    break
            
            if not is_header and current_section:
                found_sections[current_section].append(text)
                
            full_text.append(text)
            
        summary = "### 計畫書分析摘要 ###\n"
        summary += f"檔案: {os.path.basename(docx_path)}\n"
        summary += f"總段落數: {len(doc.paragraphs)}\n"
        
        if found_sections:
            summary += "\n偵測到的章節:\n"
            for section, content in found_sections.items():
                summary += f"- {section}: {len(content)} 行\n"
                for line in content[:3]:
                    summary += f"  > {line[:60]}...\n"
        else:
            summary += "\n未偵測到特定章節 (如計畫目標/KPI)。\n"
            summary += "前 500 個字元:\n"
            summary += "\n".join(full_text)[:500] + "...\n"
            
        return summary

    except Exception as e:
        return f"讀取 .docx 檔案時發生錯誤: {e}"

def extract_report_data(xlsx_path):
    """
    Extracts data from a performance report (Excel).
    """
    if not os.path.exists(xlsx_path):
        return f"錯誤: 找不到檔案: {xlsx_path}"
        
    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        sheet = wb.active
        
        summary = "### 執行報告分析摘要 ###\n"
        summary += f"檔案: {os.path.basename(xlsx_path)}\n"
        summary += f"工作表: {sheet.title}\n"
        
        # Read header
        headers = []
        for cell in sheet[1]:
            headers.append(str(cell.value) if cell.value else "欄位")
            
        summary += f"標題列: {', '.join(headers)}\n"
        
        # Preview rows
        summary += "\n前 3 筆資料預覽:\n"
        for row in sheet.iter_rows(min_row=2, max_row=4, values_only=True):
            # Convert row tuple to string for display, handle None
            row_display = [str(val) if val is not None else "" for val in row]
            summary += f"  {row_display}\n"
            
        return summary

    except Exception as e:
        return f"讀取 .xlsx 檔案時發生錯誤: {e}"

def main():
    parser = argparse.ArgumentParser(description="Phantom Policy Strategist")
    parser.add_argument("--plan", help="計畫書 .docx 路徑")
    parser.add_argument("--report", help="報告 .xlsx 路徑")
    parser.add_argument("--output", help="輸出檔案路徑 (選填)")
    
    args = parser.parse_args()
    
    output_text = ""
    
    if args.plan:
        output_text += extract_plan_text(args.plan) + "\n" + ("-" * 30) + "\n"
        
    if args.report:
        output_text += extract_report_data(args.report) + "\n" + ("-" * 30) + "\n"
        
    if args.output:
        try:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output_text)
            print(f"報告已寫入: {args.output}")
        except Exception as e:
            print(f"寫入檔案失敗: {e}")
    else:
        print(output_text)

if __name__ == "__main__":
    main()
