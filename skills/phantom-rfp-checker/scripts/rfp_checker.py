#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
rfp_checker.py - 核心檢核引擎 (V3.2.1 數據深度對核修正版)
======================================================
"""

import os
import sys
import json
import re
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
import xlsxwriter

def clean_amount(text):
    if not text: return 0.0
    text = text.replace(',', '')
    # 提取所有整數或浮點數
    nums = re.findall(r'\d+(?:\.\d+)?', text)
    if not nums: return 0.0
    try:
        return float(max(nums, key=float))
    except:
        return 0.0

def extract_kpi_targets(text):
    """強化版 KPI 提取邏輯，使用更精確的正則感應"""
    kpis = {}
    # 模式定義
    patterns = {
        '量化人次/場次': [
            r'(\d[,0-9]*)\s*(?:人次|名|場次)',
            r'(?:達|超過|至少)\s*(\d[,0-9]*)'
        ],
        '滿意度/比率': [
            r'(\d+(?:\.\d+)?)\s*(?:%|百分比)',
            r'滿意度\s*(?:達|目標為)\s*(\d+(?:\.\d+)?)\s*%'
        ]
    }
    
    for label, regex_list in patterns.items():
        found_values = []
        for p in regex_list:
            matches = re.findall(p, text)
            for m in matches:
                val = float(m.replace(',', ''))
                if val > 0: found_values.append(val)
        if found_values:
            kpis[label] = max(found_values)
            
    return kpis

def extract_budget_items(doc_path):
    budget_data = {}
    if not os.path.exists(doc_path): return budget_data
    try:
        doc = Document(doc_path)
        keywords = ['人事費', '業務費', '管理費', '設備費', '總計', '預算金額', '經費', '合計']
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2: continue
                # 掃描每一列
                for i in range(len(cells) - 1):
                    label = cells[i]
                    value = cells[i+1]
                    if any(kw in label for kw in keywords):
                        amount = clean_amount(value)
                        if amount > 0: budget_data[label] = amount
        return budget_data
    except:
        return budget_data

def get_text_segments(doc_path):
    if not os.path.exists(doc_path): return []
    try:
        doc = Document(doc_path)
        segments = []
        current_head = "DOCUMENT_START"
        current_lines = []
        def is_likely_heading(text):
            if not text or len(text) > 60: return False
            patterns = [r'^[壹貳參肆伍陸柒捌玖拾]\s*[、.]', r'^[一二三四五六七八九十]\s*[、.]', r'^\d+\s*[、.]']
            return any(re.match(p, text) for p in patterns)
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t: continue
            if is_likely_heading(t):
                if current_lines: segments.append({"header": current_head, "content": "\n".join(current_lines)})
                current_head = t
                current_lines = []
            else: current_lines.append(t)
        if current_lines: segments.append({"header": current_head, "content": "\n".join(current_lines)})
        return segments
    except: return []

def find_best_segment(segments, rule):
    name = rule['name']
    for seg in segments:
        if name in seg['header'] or any(a in seg['header'] for a in rule.get('aliases', [])):
            return seg['content'], seg['header']
    return "", "N/A"

def run_comparison(rfp_segments, plan_segments, rules, rfp_path, plan_path):
    results = []
    rfp_budget = extract_budget_items(rfp_path)
    plan_budget = extract_budget_items(plan_path)

    for sec in rules['mandatory_sections']:
        rfp_content, rfp_loc = find_best_segment(rfp_segments, sec)
        plan_content, plan_loc = find_best_segment(plan_segments, sec)
        status = "[PASS]" if rfp_content else "[FAIL]"
        details = []
        kpi_info = "N/A"
        
        # 預算邏輯
        if any(kw in sec['name'] for kw in ['經費', '預算']):
            p_total = plan_budget.get('總計', plan_budget.get('合計', 0))
            r_limit = rfp_budget.get('總計', rfp_budget.get('預算金額', 0))
            if r_limit > 0 and p_total > r_limit:
                status = "[FAIL]"
                details.append(f"預算超支：{p_total:,.0f} > 上限 {r_limit:,.0f}")
            elif r_limit > 0:
                kpi_info = f"預算對比：{p_total:,.0f} / {r_limit:,.0f}"

        # KPI 對核
        elif any(kw in sec['name'] for kw in ['KPI', '指標']):
            r_k = extract_kpi_targets(rfp_content)
            p_k = extract_kpi_targets(plan_content)
            comp = []
            for k, rv in r_k.items():
                pv = p_k.get(k, 0)
                if pv < rv:
                    status = "[FAIL]"
                    details.append(f"{k}未達標: {pv}/{rv}")
                else:
                    comp.append(f"{k}:{pv}/{rv}")
            if comp: kpi_info = " | ".join(comp)

        results.append({
            "dim": sec['name'],
            "soul": f"要點: {', '.join(sec.get('keywords', []))}",
            "body": f"[{rfp_loc}] {rfp_content[:200]}...",
            "kpi": kpi_info,
            "suggest": "; ".join(details) if details else "符合規範",
            "priority": status
        })
    return results

def generate_v3_excel(results, output_path):
    wb = xlsxwriter.Workbook(output_path)
    ws = wb.add_worksheet('Analysis')
    header_f = wb.add_format({'bold':True, 'bg_color':'#DDEAF6', 'border':1})
    cell_f = wb.add_format({'border':1, 'text_wrap':True, 'valign':'top'})
    headers = ['項次', '維度', '要點', '片段', '數據', '建議', '狀態']
    for i, h in enumerate(headers): ws.write(0, i, h, header_f)
    for i, res in enumerate(results):
        row = i + 1
        ws.write(row, 0, i+1, cell_f)
        ws.write(row, 1, res['dim'], cell_f)
        ws.write(row, 2, res['soul'], cell_f)
        ws.write(row, 3, res['body'], cell_f)
        ws.write(row, 4, res['kpi'], cell_f)
        ws.write(row, 5, res['suggest'], cell_f)
        ws.write(row, 6, res['priority'], cell_f)
    wb.close()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--rfp", required=True); parser.add_argument("--plan", required=True)
    parser.add_argument("--rules", required=True); parser.add_argument("--output", required=True)
    args = parser.parse_args()
    with open(args.rules, 'r', encoding='utf-8-sig') as f: rules = json.load(f)
    rs = run_comparison(get_text_segments(args.rfp), get_text_segments(args.plan), rules, args.rfp, args.plan)
    generate_v3_excel(rs, args.output)
    print("[OK] Complete.")

if __name__ == "__main__":
    main()


