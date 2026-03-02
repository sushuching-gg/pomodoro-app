import argparse
import os
import sys
import re
from datetime import datetime

# Ensure UTF-8 output on Windows
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)

try:
    from docx import Document
except ImportError:
    print("\u932f\u8aa4: \u672a\u5b89\u88dd python-docx\u3002\u8acb\u57f7\u884c: pip install python-docx")
    sys.exit(1)

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import fitz  # PyMuPDF
    import easyocr
    import numpy as np
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


# ============================================================
# Document Parser
# ============================================================

class DocParser:
    """Parse a .docx file and extract key sections for comparison."""

    BUDGET_KEYWORDS = ['\u7d93\u8cbb', '\u9810\u7b97', '\u55ae\u50f9', '\u91d1\u984d', '\u8cbb\u7528', '\u7d93\u8cbb\u7de8\u5217',
                       '\u6982\u7b97', '\u8cc7\u672c\u9580', 'Capital', '\u5c0f\u8a08', '\u5408\u8a08', '\u7e3d\u8a08']
    KPI_KEYWORDS = ['KPI', '\u95dc\u9375\u7e3e\u6548', '\u7e3e\u6548\u6307\u6a19', '\u9054\u6210\u7387', '\u76ee\u6a19\u503c',
                    '\u91cf\u5316\u6307\u6a19', '\u8a55\u4f30\u6307\u6a19', '\u7d44\u6210\u7d50\u69cb']
    WORK_KEYWORDS = ['\u5de5\u4f5c\u9805\u76ee', '\u57f7\u884c\u5167\u5bb9', '\u670d\u52d9\u9805\u76ee', '\u670d\u52d9\u7bc4\u570d',
                     '\u5be6\u65bd\u65b9\u6cd5', '\u57f7\u884c\u7b56\u7565', '\u670d\u52d9\u5167\u5bb9']
    SCHEDULE_KEYWORDS = ['\u7518\u7279\u5716', '\u6642\u7a0b', '\u9032\u5ea6', '\u57f7\u884c\u671f\u7a0b', '\u5de5\u4f5c\u671f\u7a0b',
                         '\u6708\u4efd', '\u5b63\u5ea6', '\u91cc\u7a0b\u7891']
    TEAM_KEYWORDS = ['\u4eba\u529b', '\u5718\u968a', '\u4eba\u54e1', '\u914d\u7f6e', '\u7d44\u7e54\u67b6\u69cb',
                     '\u5c08\u6848\u7d93\u7406', '\u5c08\u6848\u4eba\u54e1']

    def __init__(self, path, doc_type):
        self.path = path
        self.doc_type = doc_type  # 'plan', 'rfp', 'proposal'
        self.doc = None
        self.data = {
            'filename': os.path.basename(path),
            'city': None,
            'year': None,
            'budget_tables': [],
            'kpi_sections': [],
            'work_items': [],
            'schedule_info': [],
            'team_info': [],
            'full_text': '',
            'total_paragraphs': 0,
            'total_tables': 0,
        }

    def parse(self):
        if not os.path.exists(self.path):
            print(f"\u932f\u8aa4: \u627e\u4e0d\u5230\u6a94\u6848 - {self.path}")
            return self.data

        ext = os.path.splitext(self.path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf()
        elif ext in ['.docx', '.doc']:
            return self._parse_docx()
        else:
            print(f"\u932f\u8aa4: \u4e0d\u652f\u63f4\u7684\u683c\u5f0f: {ext}")
            return self.data
    
    def _parse_pdf(self):
        if not HAS_PDF:
            print("\u932f\u8aa4: \u672a\u5b89\u88dd PyPDF2\u3002\u8acb\u57f7\u884c: pip install PyPDF2")
            return self.data
        
        try:
            reader = PdfReader(self.path)
            all_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
            
            full_text = '\n'.join(all_text)
            self.data['total_paragraphs'] = full_text.count('\n')
            self.data['total_tables'] = 0  # PDF table detection is limited
            self.data['full_text'] = full_text
            
            self._extract_metadata()
            self._parse_text_content(full_text)
            self._extract_budget_from_text(full_text)
            
        except Exception as e:
            print(f"\u932f\u8aa4: \u7121\u6cd5\u8b80\u53d6 PDF {self.path}: {e}")
        
        return self.data
    
    def _parse_text_content(self, full_text):
        """Parse sections from plain text (used for PDF)."""
        lines = full_text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
            
            if any(kw in text for kw in self.KPI_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'kpi'
                section_content = [text]
            elif any(kw in text for kw in self.WORK_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'work'
                section_content = [text]
            elif any(kw in text for kw in self.SCHEDULE_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'schedule'
                section_content = [text]
            elif any(kw in text for kw in self.TEAM_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'team'
                section_content = [text]
            elif current_section:
                section_content.append(text)
        
        if current_section and section_content:
            self._store_section(current_section, section_content)
    
    def _extract_budget_from_text(self, text):
        """Try to extract budget info from PDF text."""
        lines = text.split('\n')
        in_budget = False
        budget_rows = []
        
        for line in lines:
            stripped = line.strip()
            if any(kw in stripped for kw in self.BUDGET_KEYWORDS):
                in_budget = True
                budget_rows.append([stripped])
            elif in_budget:
                if stripped and len(stripped) > 2:
                    budget_rows.append([stripped])
                if len(budget_rows) > 20:
                    break
                if not stripped:
                    if budget_rows:
                        break
        
        if budget_rows:
            self.data['budget_tables'].append({
                'index': 0,
                'rows': budget_rows,
                'row_count': len(budget_rows),
                'source': 'pdf_text'
            })

    def _parse_docx(self):
        try:
            self.doc = Document(self.path)
        except Exception as e:
            print(f"\u932f\u8aa4: \u7121\u6cd5\u8b80\u53d6 {self.path}: {e}")
            return self.data

        self.data['total_paragraphs'] = len(self.doc.paragraphs)
        self.data['total_tables'] = len(self.doc.tables)

        self._extract_metadata()
        self._extract_text_sections()
        self._extract_tables()

        return self.data

    def _extract_metadata(self):
        filename = self.data['filename']
        full_path = self.path

        # City detection
        cities = ['\u81fa\u5317\u5e02','\u65b0\u5317\u5e02','\u6843\u5712\u5e02','\u81fa\u4e2d\u5e02','\u81fa\u5357\u5e02','\u9ad8\u96c4\u5e02',
                  '\u57fa\u9686\u5e02','\u65b0\u7af9\u5e02','\u65b0\u7af9\u7e23','\u82d7\u6817\u7e23','\u5f70\u5316\u7e23','\u5357\u6295\u7e23',
                  '\u96f2\u6797\u7e23','\u5609\u7fa9\u7e23','\u5609\u7fa9\u5e02','\u5c4f\u6771\u7e23','\u5b9c\u862d\u7e23','\u82b1\u84ee\u7e23',
                  '\u81fa\u6771\u7e23','\u6f8e\u6e56\u7e23','\u91d1\u9580\u7e23','\u9023\u6c5f\u7e23']
        for city in cities:
            if city in filename or city in full_path:
                self.data['city'] = city
                break

        # Year detection
        year_match = re.search(r'(1[0-2][0-9])\u5e74', filename)
        if year_match:
            self.data['year'] = year_match.group(1)
        else:
            year_match2 = re.search(r'(1[0-2][0-9])', filename)
            if year_match2:
                self.data['year'] = year_match2.group(1)

    def _extract_text_sections(self):
        current_section = None
        section_content = []
        all_text = []

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)

            # Detect section headers
            if any(kw in text for kw in self.KPI_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'kpi'
                section_content = [text]
            elif any(kw in text for kw in self.WORK_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'work'
                section_content = [text]
            elif any(kw in text for kw in self.SCHEDULE_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'schedule'
                section_content = [text]
            elif any(kw in text for kw in self.TEAM_KEYWORDS):
                if current_section and section_content:
                    self._store_section(current_section, section_content)
                current_section = 'team'
                section_content = [text]
            elif current_section:
                section_content.append(text)

        # Store last section
        if current_section and section_content:
            self._store_section(current_section, section_content)

        self.data['full_text'] = '\n'.join(all_text)

    def _store_section(self, section_type, content):
        target = {
            'kpi': 'kpi_sections',
            'work': 'work_items',
            'schedule': 'schedule_info',
            'team': 'team_info',
        }.get(section_type)
        if target:
            self.data[target].extend(content[:20])  # Cap at 20 lines

    def _extract_tables(self):
        for i, table in enumerate(self.doc.tables):
            rows = []
            is_budget = False
            try:
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    if any(cells):
                        rows.append(cells)
                # Check if budget table
                if rows:
                    header_str = ' '.join(rows[0])
                    if any(kw in header_str for kw in self.BUDGET_KEYWORDS):
                        is_budget = True
            except Exception:
                continue

            if is_budget and rows:
                self.data['budget_tables'].append({
                    'index': i,
                    'rows': rows,
                    'row_count': len(rows)
                })


# ============================================================
# Comparator
# ============================================================

class TripleDocComparator:
    DOC_LABELS = {
        'plan': '\u6838\u5b9a\u8a08\u756b\u66f8',
        'rfp': '\u62db\u6a19\u9700\u6c42\u66f8',
        'proposal': '\u670d\u52d9\u5efa\u8b70\u66f8',
    }

    def __init__(self, plan_data, rfp_data, proposal_data):
        self.docs = {
            'plan': plan_data,
            'rfp': rfp_data,
            'proposal': proposal_data,
        }
        self.diffs = []

    def compare_all(self):
        report = []
        city = self._get_city()
        year = self._get_year()

        report.append(f"# \u4e09\u66f8\u6bd4\u5c0d\u5dee\u7570\u5c0d\u7167\u8868")
        report.append(f"\n> \u7522\u751f\u6642\u9593: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if city: report.append(f"> \u7e23\u5e02: {city}")
        if year: report.append(f"> \u5e74\u5ea6: {year}")
        report.append("")

        # 1. Document Overview
        report.append("## 1. \u6587\u4ef6\u6982\u89bd")
        report.append("")
        report.append("| \u9805\u76ee | \u6838\u5b9a\u8a08\u756b\u66f8 | \u62db\u6a19\u9700\u6c42\u66f8 | \u670d\u52d9\u5efa\u8b70\u66f8 |")
        report.append("|------|------------|------------|------------|")
        report.append(f"| \u6a94\u540d | {self.docs['plan']['filename']} | {self.docs['rfp']['filename']} | {self.docs['proposal']['filename']} |")
        report.append(f"| \u6bb5\u843d\u6578 | {self.docs['plan']['total_paragraphs']} | {self.docs['rfp']['total_paragraphs']} | {self.docs['proposal']['total_paragraphs']} |")
        report.append(f"| \u8868\u683c\u6578 | {self.docs['plan']['total_tables']} | {self.docs['rfp']['total_tables']} | {self.docs['proposal']['total_tables']} |")
        budget_counts = [len(d['budget_tables']) for d in self.docs.values()]
        report.append(f"| \u9810\u7b97\u8868 | {budget_counts[0]} | {budget_counts[1]} | {budget_counts[2]} |")
        report.append("")

        # 2. Budget Comparison
        report.append("## 2. \u7d93\u8cbb\u7de8\u5217\u6bd4\u5c0d")
        report.append("")
        budget_report = self._compare_budgets()
        report.extend(budget_report)

        # 3. Work Items
        report.append("## 3. \u5de5\u4f5c\u9805\u76ee\u6bd4\u5c0d")
        report.append("")
        work_report = self._compare_sections('work_items', '\u5de5\u4f5c\u9805\u76ee')
        report.extend(work_report)

        # 4. KPI
        report.append("## 4. KPI/\u7e3e\u6548\u6307\u6a19\u6bd4\u5c0d")
        report.append("")
        kpi_report = self._compare_sections('kpi_sections', 'KPI')
        report.extend(kpi_report)

        # 5. Schedule
        report.append("## 5. \u6642\u7a0b/\u7518\u7279\u5716\u6bd4\u5c0d")
        report.append("")
        schedule_report = self._compare_sections('schedule_info', '\u6642\u7a0b')
        report.extend(schedule_report)

        # 6. Team
        report.append("## 6. \u4eba\u529b\u914d\u7f6e\u6bd4\u5c0d")
        report.append("")
        team_report = self._compare_sections('team_info', '\u4eba\u529b')
        report.extend(team_report)

        # 7. Risk Summary
        report.append("## 7. \u98a8\u96aa\u63d0\u793a\u8207\u5efa\u8b70")
        report.append("")
        risk_report = self._generate_risk_summary()
        report.extend(risk_report)

        return '\n'.join(report)

    def _get_city(self):
        for d in self.docs.values():
            if d.get('city'):
                return d['city']
        return None

    def _get_year(self):
        for d in self.docs.values():
            if d.get('year'):
                return d['year']
        return None

    def _compare_budgets(self):
        lines = []
        has_any = False

        for doc_type, label in self.DOC_LABELS.items():
            tables = self.docs[doc_type]['budget_tables']
            if tables:
                has_any = True
                for bt in tables:
                    lines.append(f"### {label} - \u9810\u7b97\u8868 (Table #{bt['index']+1}, {bt['row_count']} \u5217)")
                    lines.append("")
                    # Show first row (header) and last row (total)
                    if bt['rows']:
                        header = ' | '.join(bt['rows'][0][:5])
                        lines.append(f"- \u6a19\u984c\u5217: {header}")
                        if len(bt['rows']) > 1:
                            last = ' | '.join(bt['rows'][-1][:5])
                            lines.append(f"- \u6700\u5f8c\u4e00\u5217 (\u5408\u8a08): {last}")
                    lines.append("")

        if not has_any:
            lines.append("\u26a0 \u4e09\u4efd\u6587\u4ef6\u4e2d\u5747\u672a\u5075\u6e2c\u5230\u9810\u7b97\u8868\u3002\u8acb\u4eba\u5de5\u78ba\u8a8d\u7d93\u8cbb\u7de8\u5217\u662f\u5426\u4e00\u81f4\u3002")
            lines.append("")
            self.diffs.append('\u7f3a\u5c11\u9810\u7b97\u8868')

        # Cross-check: do all docs have budget tables?
        missing_budget = [self.DOC_LABELS[k] for k, v in self.docs.items() if not v['budget_tables']]
        if missing_budget and has_any:
            lines.append(f"\u26a0 \u4ee5\u4e0b\u6587\u4ef6\u672a\u5075\u6e2c\u5230\u9810\u7b97\u8868: **{', '.join(missing_budget)}**")
            lines.append("\u8acb\u78ba\u8a8d\u662f\u5426\u5728\u9644\u4ef6\u4e2d\u6216\u4f7f\u7528\u4e0d\u540c\u683c\u5f0f\u3002")
            lines.append("")
            self.diffs.append(f"\u9810\u7b97\u8868\u7f3a\u5931: {', '.join(missing_budget)}")

        return lines

    def _compare_sections(self, section_key, section_name):
        lines = []

        for doc_type, label in self.DOC_LABELS.items():
            content = self.docs[doc_type].get(section_key, [])
            if content:
                lines.append(f"### {label}")
                for item in content[:8]:
                    lines.append(f"- {item[:80]}")
                if len(content) > 8:
                    lines.append(f"- ... (\u5171 {len(content)} \u9805)")
                lines.append("")
            else:
                lines.append(f"### {label}")
                lines.append(f"\u26a0 \u672a\u5075\u6e2c\u5230{section_name}\u76f8\u95dc\u5167\u5bb9\u3002")
                lines.append("")

        # Cross-check
        has_content = {k: bool(self.docs[k].get(section_key, [])) for k in self.docs}
        if any(has_content.values()) and not all(has_content.values()):
            missing = [self.DOC_LABELS[k] for k, v in has_content.items() if not v]
            lines.append(f"\u26a0 **\u5dee\u7570**: {section_name}\u5728\u4ee5\u4e0b\u6587\u4ef6\u4e2d\u7f3a\u5931: {', '.join(missing)}")
            lines.append("")
            self.diffs.append(f"{section_name}\u7f3a\u5931: {', '.join(missing)}")

        return lines

    def _generate_risk_summary(self):
        lines = []

        if not self.diffs:
            lines.append("\u2713 \u672a\u767c\u73fe\u660e\u986f\u7d50\u69cb\u6027\u5dee\u7570\u3002\u5efa\u8b70\u4ecd\u9700\u4eba\u5de5\u6838\u5c0d\u7d30\u7bc0\u5167\u5bb9\u3002")
        else:
            lines.append(f"\u26a0 \u5171\u767c\u73fe **{len(self.diffs)}** \u9805\u6f5b\u5728\u5dee\u7570:")
            lines.append("")
            for i, diff in enumerate(self.diffs, 1):
                lines.append(f"{i}. {diff}")
            lines.append("")
            lines.append("\u5efa\u8b70:")
            lines.append("- \u8acb\u91dd\u5c0d\u4ee5\u4e0a\u5dee\u7570\u9805\u76ee\u9032\u884c\u4eba\u5de5\u6838\u5c0d")
            lines.append("- \u7d93\u8cbb\u7de8\u5217\u5dee\u7570\u61c9\u5728\u8a2a\u8996\u524d\u8207\u7e23\u5e02\u78ba\u8a8d")
            lines.append("- KPI \u4e0d\u4e00\u81f4\u53ef\u80fd\u5f71\u97ff\u7d50\u6848\u8a55\u4f30")

        lines.append("")
        lines.append("---")
        lines.append("*\u6b64\u5831\u544a\u7531 Phantom \u4e09\u66f8\u6bd4\u5c0d Agent \u81ea\u52d5\u7522\u751f\uff0c\u50c5\u4f9b\u53c3\u8003\uff0c\u6b63\u5f0f\u6bd4\u5c0d\u4ecd\u9700\u4eba\u5de5\u78ba\u8a8d\u3002*")
        return lines


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Phantom \u4e09\u66f8\u81ea\u52d5\u6bd4\u5c0d (Triple-Doc Checker)")
    parser.add_argument("--plan", required=True, help="\u6838\u5b9a\u8a08\u756b\u66f8\u8def\u5f91 (.docx)")
    parser.add_argument("--rfp", required=True, help="\u62db\u6a19\u9700\u6c42\u66f8\u8def\u5f91 (.docx)")
    parser.add_argument("--proposal", required=True, help="\u670d\u52d9\u5efa\u8b70\u66f8\u8def\u5f91 (.docx)")
    parser.add_argument("--output", default=None, help="\u8f38\u51fa\u6a94\u6848\u8def\u5f91 (\u9078\u586b\uff0c\u9810\u8a2d\u8f38\u51fa\u81f3\u7d42\u7aef)")

    args = parser.parse_args()

    print("=" * 50)
    print("  Phantom \u4e09\u66f8\u81ea\u52d5\u6bd4\u5c0d (Triple-Doc Checker)")
    print("=" * 50)
    print()

    # Parse all three documents
    print("[\u6b65\u9a5f 1/3] \u89e3\u6790\u6838\u5b9a\u8a08\u756b\u66f8...")
    plan_parser = DocParser(args.plan, 'plan')
    plan_data = plan_parser.parse()

    print("[\u6b65\u9a5f 2/3] \u89e3\u6790\u62db\u6a19\u9700\u6c42\u66f8...")
    rfp_parser = DocParser(args.rfp, 'rfp')
    rfp_data = rfp_parser.parse()

    print("[\u6b65\u9a5f 3/3] \u89e3\u6790\u670d\u52d9\u5efa\u8b70\u66f8...")
    proposal_parser = DocParser(args.proposal, 'proposal')
    proposal_data = proposal_parser.parse()

    # Compare
    print("\n[\u6bd4\u5c0d\u4e2d] \u6b63\u5728\u6bd4\u5c0d\u4e09\u66f8\u5dee\u7570...\n")
    comparator = TripleDocComparator(plan_data, rfp_data, proposal_data)
    report = comparator.compare_all()

    # Output
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"\u2713 \u6bd4\u5c0d\u5831\u544a\u5df2\u5132\u5b58\u81f3: {args.output}")
    else:
        print(report)

    print("\n\u6bd4\u5c0d\u5b8c\u6210\u3002")


if __name__ == "__main__":
    main()
