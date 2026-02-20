import sys
import os
import json
import re
import docx

def parse_plan(docx_path):
    print(f'Parsing: {docx_path}')
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        return {'error': str(e)}
    
    filename = os.path.basename(docx_path)
    parent_dir = os.path.basename(os.path.dirname(docx_path))
    
    data = {
        'file_path': docx_path,
        'filename': filename,
        'metadata': {},
        'doc_type': 'proposal',
        'content_summary': '',
        'tables': []
    }

    if any(kw in filename for kw in ['審查', '意見', '彙整', '回覆']):
        data['doc_type'] = 'review_form'

    # 1. Filename Parsing (Date priority)
    # yyyMMdd (e.g. 1130520)
    date_match = re.search(r'(1[0-2][0-9])([0-1][0-9])([0-3][0-9])', filename)
    if date_match and int(date_match.group(2)) <= 12:
        data['metadata']['year'] = date_match.group(1)
        data['metadata']['month'] = str(int(date_match.group(2)))
    
    if 'year' not in data['metadata']:
        year_match = re.search(r'(1[0-2][0-9])年', filename)
        if year_match:
            data['metadata']['year'] = year_match.group(1)
        else:
            ad_year_match = re.search(r'(20[1-3][0-9])', filename)
            if ad_year_match:
                ad_year = int(ad_year_match.group(1))
                data['metadata']['year'] = str(ad_year - 1911)

    # 2. Deep Content Parsing
    try:
        for para in doc.paragraphs[:50]:
            text = para.text.strip()
            if not text: continue
            
            # Year Extraction
            if 'year' not in data['metadata']:
                roc_match = re.search(r'中華民國\s*([1-9][0-9]{1,2})\s*年', text)
                if roc_match:
                    data['metadata']['year'] = roc_match.group(1)
                
                if 'year' not in data['metadata']:
                    y_match = re.search(r'([1-9][0-9]{1,2})\s*年度', text)
                    if y_match:
                        data['metadata']['year'] = y_match.group(1)
                
                if 'year' not in data['metadata']:
                    ad_match = re.search(r'(20[1-3][0-9])年', text)
                    if ad_match:
                        data['metadata']['year'] = str(int(ad_match.group(1)) - 1911)

            # Month Extraction
            if 'month' not in data['metadata']:
                # 10月, 5月
                m_match = re.search(r'年\s*([0-1]?[0-9])\s*月', text)
                if m_match:
                    try:
                        m = int(m_match.group(1))
                        if 1 <= m <= 12:
                            data['metadata']['month'] = str(m)
                    except: pass
                
                if 'month' not in data['metadata']:
                    # Chinese Month
                    cn_months = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, '七':7, '八':8, '九':9, '十':10, '十一':11, '十二':12}
                    for k,v in cn_months.items():
                        if f'{k}月' in text:
                            data['metadata']['month'] = str(v)
                            break
                            
            if 'year' in data['metadata'] and 'month' in data['metadata']:
                break
    except Exception:
        pass

    # City Extraction
    cities = ['臺北市', '新北市', '桃園市', '臺中市', '臺南市', '高雄市', '基隆市', '新竹市', '新竹縣', '苗栗縣', '彰化縣', '南投縣', '雲林縣', '嘉義縣', '嘉義市', '屏東縣', '宜蘭縣', '花蓮縣', '臺東縣', '澎湖縣', '金門縣', '連江縣']
    
    found_city = None
    for city in cities:
        if city in filename:
            found_city = city
            break
            
    if not found_city:
        doc_str = str(docx_path)
        for city in cities:
            if city in doc_str:
                found_city = city
                break

    if found_city:
        data['metadata']['city'] = found_city

    # Summary
    summary_paras = []
    if data['doc_type'] == 'proposal':
        for para in doc.paragraphs[:30]:
            text = para.text.strip()
            if len(text) > 10:
                summary_paras.append(text)
        data['content_summary'] = '\\n'.join(summary_paras[:6])
    else:
        data['content_summary'] = '(Review Form - Content Skipped)'

    # Tables
    if data['doc_type'] == 'proposal':
        for i, table in enumerate(doc.tables):
            rows = []
            is_budget_table = False
            if table.rows:
                try:
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    header_str = ''.join(header_cells)
                    budget_keywords = ['經費', '預算', '單價', '金額', '費用', 'Capital']
                    if any(kw in header_str for kw in budget_keywords):
                        is_budget_table = True
                    
                    if is_budget_table:
                        for row in table.rows:
                            cells = [cell.text.strip().replace('\\n', ' ') for cell in row.cells]
                            if any(cells):
                                rows.append(cells)
                except Exception:
                    pass
            if is_budget_table:
                data['tables'].append({
                    'type': 'budget_candidate',
                    'index': i,
                    'content': rows
                })
            
    return data # end function

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python plan_parser.py <docx_path>')
        sys.exit(1)
    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        sys.exit(1)
    result = parse_plan(docx_path)
    output_filename = 'extracted_plan.json'
    with open(output_filename, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f'Extraction complete. Saved to {output_filename}')
